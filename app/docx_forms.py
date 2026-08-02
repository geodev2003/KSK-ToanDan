# -*- coding: utf-8 -*-
"""
Sinh PHIẾU KHÁM SỨC KHỎE (.docx) đúng cấu trúc "Mẫu 02" (Thông tư 25/2026/TT-BYT)
mà bệnh viện đang dùng. Phần THÔNG TIN HÀNH CHÍNH lấy tự động từ dữ liệu đã nhập
trong app; các phần còn lại (tiền sử, tiêm chủng, khám lâm sàng, kết luận...) để
trống theo đúng bố cục gốc cho người khai / bác sĩ điền tay khi khám.

File PDF được xuất bằng cách convert chính file .docx này qua LibreOffice headless
(xem hàm docx_to_pdf) — nên bản Word và bản PDF LUÔN khớp định dạng với nhau.
"""
import os
import subprocess
import tempfile
import uuid

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Times New Roman"
CHECK = "☐"
CHECKED = "☒"


# ---------------- helpers định dạng ----------------
def _set_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def _p(doc, text="", size=12, bold=False, italic=False, align=None, space_after=4, space_before=0):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    run = para.add_run(text)
    _set_font(run, size, bold, italic)
    return para


def _field_line(doc, label, value, dots=60):
    """Dòng dạng '1. Nhãn: giá trị..........' — nếu có value thì điền, không thì để chấm."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    r1 = para.add_run(label + ": ")
    _set_font(r1, 12, bold=False)
    if value:
        r2 = para.add_run(str(value))
        _set_font(r2, 12, bold=True)
        r3 = para.add_run(" " + "." * max(0, dots - len(str(value))))
        _set_font(r3, 12)
    else:
        r2 = para.add_run("." * dots)
        _set_font(r2, 12)
    return para


def _checkbox_line(doc, options, prefix=""):
    """options: list[(label, checked_bool)]"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    if prefix:
        r_pre = para.add_run(prefix)
        _set_font(r_pre, 12, bold=False)
    for i, (label, checked) in enumerate(options):
        r = para.add_run(("     " if (i or prefix) else "") + (CHECKED if checked else CHECK) + " " + label)
        _set_font(r, 12)
    return para


def _heading(doc, text, size=13):
    return _p(doc, text, size=size, bold=True, space_before=10, space_after=6)


def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_text(cell, text, size=10, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    _set_font(r, size, bold)


def _signature_block(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("………, ngày ….. tháng ….. năm ……")
    _set_font(r, 11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(title)
    _set_font(r2, 11, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p3.add_run("(Ký, ghi rõ họ tên" + (" và đóng dấu)" if "KẾT LUẬN" in title else ")"))
    _set_font(r3, 10.5, italic=True)
    for _ in range(2):
        doc.add_paragraph()  # chừa khoảng trống để ký tay


def _new_table(doc, rows, cols, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


# ---------------- các bảng chuẩn theo Mẫu 02 ----------------
VACCINE_ROWS = [
    ("Lao", "Sơ sinh", "Mũi 1"),
    ("Viêm gan siêu vi B", "Trong 24h sau sinh", "Mũi 1"),
    ("", "02 tháng", "Mũi 2"),
    ("", "03 tháng", "Mũi 3"),
    ("", "04 tháng", "Mũi 4"),
    ("Bạch hầu - Uốn ván - Ho gà", "02 tháng", "Mũi 1"),
    ("", "03 tháng", "Mũi 2"),
    ("", "04 tháng", "Mũi 3"),
    ("", "18 tháng", "Mũi 4"),
    ("Bại liệt", "02-03-04 tháng (uống)", "Liều 1-3"),
    ("", "05, 09 tháng (tiêm)", "Mũi 1-2"),
    ("Hib (viêm phổi/màng não)", "02-03-04 tháng", "Mũi 1-3"),
    ("Sởi", "09 / 18 tháng", "Mũi 1-2"),
    ("Sởi - Rubella", "18 tháng", "Mũi 1"),
    ("Viêm não Nhật Bản", "12 / +1-2 tuần / 24 tháng", "Mũi 1-3"),
    ("Khác", "…………………", "…………"),
]

CLINICAL_SPECIALTIES = [
    ("1.", "Nhi khoa", [
        ("a)", "Tuần hoàn"), ("b)", "Hô hấp"), ("c)", "Tiêu hóa"),
        ("d)", "Thận - Tiết niệu"), ("đ)", "Thần kinh"), ("e)", "Tâm thần"),
        ("g)", "Khám lâm sàng khác"),
    ]),
]
OTHER_SPECIALTIES = [
    ("2.", "Mắt"),
    ("3.", "Tai - Mũi - Họng"),
    ("4.", "Răng - Hàm - Mặt"),
]

LOAI = ["I", "II", "III", "IV", "V"]


def _checkbox_run(para, text, checked=False, size=10.5):
    r = para.add_run((CHECKED if checked else CHECK) + " " + text)
    _set_font(r, size)


def _add_boxes_line(doc, label, val_str, max_len=12, label_cm=8.5):
    val_str = str(val_str or "").strip()
    digits = list(val_str[:max_len]) + [""] * max(0, max_len - len(val_str))

    t = doc.add_table(rows=1, cols=1 + max_len)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False

    row = t.rows[0]
    cell_lbl = row.cells[0]
    # Set qua table.columns[i].width để cập nhật đúng <w:tblGrid>, tránh bị
    # bẻ dòng nhãn do bảng dùng độ rộng cột mặc định khi render.
    t.columns[0].width = Cm(label_cm)
    cell_lbl.width = Cm(label_cm)
    _set_cell_text(cell_lbl, label + ": ", size=11, bold=True)
    cell_lbl.paragraphs[0].paragraph_format.space_before = Pt(1)
    cell_lbl.paragraphs[0].paragraph_format.space_after = Pt(1)

    # Xóa viền cho ô chứa nhãn tên mục
    tcPr0 = cell_lbl._tc.get_or_add_tcPr()
    tcBorders0 = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders0.append(b)
    tcPr0.append(tcBorders0)

    # Thêm viền cho các ô chứa từng chữ số (ô 1 đến max_len)
    for i, d in enumerate(digits):
        cell = row.cells[1 + i]
        t.columns[1 + i].width = Cm(0.42)
        cell.width = Cm(0.42)
        _set_cell_text(cell, d, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.paragraphs[0].paragraph_format.space_before = Pt(1)
        cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            tcBorders.append(b)
        tcPr.append(tcBorders)


def _find_template_path():
    p1 = os.path.join(os.path.dirname(__file__), "templates", "Mau02.docx")
    p2 = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Mau02.docx")
    if os.path.exists(p1):
        return p1
    if os.path.exists(p2):
        return p2
    return None


def _insert_boxes_row(doc, p, label, val_str, max_len=12, label_cm=8.5, space_after_pt=6):
    parent = p._element.getparent()
    p_idx = parent.index(p._element)

    t = doc.add_table(rows=1, cols=1 + max_len)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False

    row = t.rows[0]
    cell_lbl = row.cells[0]
    # QUAN TRỌNG: phải set độ rộng qua table.columns[i].width (không chỉ
    # cell.width) để cập nhật đúng phần tử <w:tblGrid>. Nếu chỉ set
    # cell.width, Word/LibreOffice vẫn dùng độ rộng mặc định (chia đều) của
    # tblGrid khi render bảng ở chế độ layout "fixed" -> nhãn bị bẻ dòng lộn
    # xộn và các ô số không thẳng hàng với các mục còn lại của form.
    t.columns[0].width = Cm(label_cm)
    cell_lbl.width = Cm(label_cm)
    _set_cell_text(cell_lbl, label + ": ", size=11, bold=True)
    cell_lbl.paragraphs[0].paragraph_format.space_before = Pt(1)
    cell_lbl.paragraphs[0].paragraph_format.space_after = Pt(1)

    tcPr0 = cell_lbl._tc.get_or_add_tcPr()
    tcBorders0 = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders0.append(b)
    tcPr0.append(tcBorders0)

    val_str = str(val_str or "").strip()
    digits = list(val_str[:max_len]) + [""] * max(0, max_len - len(val_str))

    for i, d in enumerate(digits):
        cell = row.cells[1 + i]
        t.columns[1 + i].width = Cm(0.42)
        cell.width = Cm(0.42)
        _set_cell_text(cell, d, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.paragraphs[0].paragraph_format.space_before = Pt(1)
        cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    parent.insert(p_idx, t._element)

    # Chèn thêm 1 đoạn trống nhỏ ngay sau bảng để tạo khoảng cách dãn dòng
    # với dòng/bảng tiếp theo (bảng không có thuộc tính "space after" như
    # paragraph nên nếu không có đoạn đệm này, bảng mục 6 và mục 7 sẽ dính
    # sát vào nhau và dính sát luôn vào dòng "8. Nơi ở hiện tại").
    spacer = OxmlElement('w:p')
    spacer_pPr = OxmlElement('w:pPr')
    spacer_spacing = OxmlElement('w:spacing')
    spacer_spacing.set(qn('w:after'), str(int(space_after_pt * 20)))
    spacer_spacing.set(qn('w:line'), '20')
    spacer_spacing.set(qn('w:lineRule'), 'exact')
    spacer_pPr.append(spacer_spacing)
    spacer_rPr = OxmlElement('w:rPr')
    spacer_sz = OxmlElement('w:sz')
    spacer_sz.set(qn('w:val'), '2')
    spacer_rPr.append(spacer_sz)
    spacer_pPr.append(spacer_rPr)
    spacer.append(spacer_pPr)
    parent.insert(p_idx + 1, spacer)

    parent.remove(p._element)


def build_patient_form(rec: dict, group: dict, hospital_name: str = "BỆNH VIỆN HỒNG ĐỨC II") -> "Document":
    tpl_path = _find_template_path()
    if not tpl_path:
        return _build_form_scratch(rec, group, hospital_name)

    d = Document(tpl_path)
    gioi = (rec.get("gioi_tinh") or "").strip()
    ten_truong = group.get("ten_doan") or ""
    dia_diem = group.get("dia_diem") or ""
    lop_val = rec.get("lop") or group.get("lop") or ""

    # Mẫu 02 có 2 dòng "Xã/Phường:" liên tiếp nhau: dòng đầu thuộc mục 8
    # (Nơi ở hiện tại của trẻ), dòng thứ hai thuộc mục 11 (Địa chỉ trường).
    # Dùng cờ này để chỉ điền đúng 1 lần cho dòng đầu, tránh ghi đè nhầm dữ
    # liệu "phường" của nhà lên cả dòng "phường" của trường.
    xa_phuong_filled = False

    for p in list(d.paragraphs):
        pPr = p._p.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            numId_el = numPr.find(qn('w:numId'))
            if numId_el is not None and numId_el.attrib.get(qn('w:val')) == '1':
                pPr.remove(numPr)

        txt = p.text.strip()
        if "Mẫu 2. MẪU GIẤY KHÁM SỨC KHỎE" in txt:
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run("Mẫu 2. MẪU GIẤY KHÁM SỨC KHỎE VÀ KHÁM SỨC KHỎE ĐỊNH KỲ DÙNG CHO TRẺ TỪ ĐỦ 06 TUỔI ĐẾN 18 TUỔI")
            _set_font(r, 12.5, bold=True)
        elif "Họ và tên (viết chữ in hoa)" in txt:
            p.text = ""
            r1 = p.add_run("1. Họ và tên (viết chữ in hoa): ")
            _set_font(r1, 12)
            r2 = p.add_run((rec.get("ho_ten") or "").upper())
            _set_font(r2, 12, bold=True)
        elif "Giới tính:" in txt:
            p.text = ""
            r1 = p.add_run("2. Giới tính:  ")
            _set_font(r1, 12)
            nam_box = CHECKED if gioi == "Nam" else CHECK
            nu_box = CHECKED if gioi == "Nữ" else CHECK
            r2 = p.add_run(f"{nam_box} Nam     {nu_box} Nữ")
            _set_font(r2, 12)
        elif "Ngày tháng năm sinh:" in txt:
            p.text = ""
            r1 = p.add_run("3. Ngày tháng năm sinh: ")
            _set_font(r1, 12)
            r2 = p.add_run(str(rec.get("ngay_sinh") or ""))
            _set_font(r2, 12, bold=True)
        elif "Dân tộc:" in txt:
            p.text = ""
            r1 = p.add_run("4. Dân tộc: ")
            _set_font(r1, 12)
            r2 = p.add_run(str(rec.get("dan_toc") or ""))
            _set_font(r2, 12, bold=True)
        elif "Nhóm máu (nếu có):" in txt:
            p.text = ""
            r1 = p.add_run("5. Nhóm máu (nếu có): ")
            _set_font(r1, 12)
            nhom_mau = rec.get("nhom_mau") or ""
            if nhom_mau:
                r2 = p.add_run(str(nhom_mau))
                _set_font(r2, 12, bold=True)
            else:
                r2 = p.add_run("." * 60)
                _set_font(r2, 12)
        elif "Số CCCD/Mã số định danh/Hộ chiếu" in txt:
            _insert_boxes_row(d, p, "6. Số CCCD/Mã số định danh/Hộ chiếu", rec.get("cccd"), 12, 8.5)
        elif "Số thẻ BHYT:" in txt:
            _insert_boxes_row(d, p, "7. Số thẻ BHYT", rec.get("ma_bhyt"), 15, 8.5)
        elif "Nơi ở hiện tại:" in txt:
            p.text = ""
            r1 = p.add_run("8. Nơi ở hiện tại: ")
            _set_font(r1, 12)
            dia_chi = ", ".join(x for x in [rec.get("so_nha"), rec.get("khu_pho"), rec.get("phuong"), rec.get("tinh")] if x)
            r2 = p.add_run(dia_chi)
            _set_font(r2, 12, bold=True)
        elif "Xã/Phường:" in txt and not xa_phuong_filled:
            # Dòng Xã/Phường đầu tiên -> thuộc mục 8 (nơi ở hiện tại)
            p.text = ""
            r1 = p.add_run("Xã/Phường: ")
            _set_font(r1, 12)
            r2 = p.add_run(str(rec.get("phuong") or ""))
            _set_font(r2, 12, bold=True)
            xa_phuong_filled = True
        elif "Trẻ có đi học:" in txt:
            p.text = ""
            # Giữ ký tự tab (\t) để bám đúng các tab-stop đã định nghĩa sẵn
            # trong đoạn văn gốc, đảm bảo 2 ô checkbox thẳng cột như bản mẫu.
            r1 = p.add_run(f"9. Trẻ có đi học:\t{CHECKED} Có\t{CHECK} Không (chuyển qua câu 12)")
            _set_font(r1, 12)
        elif "Tên Trường (nếu có):" in txt:
            p.text = ""
            r1 = p.add_run("10. Tên Trường (nếu có): ")
            _set_font(r1, 12)
            r2 = p.add_run(f"{ten_truong}")
            _set_font(r2, 12, bold=True)
            r3 = p.add_run("     Lớp: ")
            _set_font(r3, 12)
            if lop_val:
                r4 = p.add_run(str(lop_val))
                _set_font(r4, 12, bold=True)
            else:
                r4 = p.add_run("." * 20)
                _set_font(r4, 12)
        elif "Địa chỉ trường:" in txt:
            p.text = ""
            r1 = p.add_run("11. Địa chỉ trường: ")
            _set_font(r1, 12)
            r2 = p.add_run(f"{dia_diem}")
            _set_font(r2, 12, bold=True)
        elif "Họ và tên mẹ hoặc người giám hộ" in txt:
            p.text = ""
            r1 = p.add_run("12. Họ và tên mẹ hoặc người giám hộ (đối với trẻ ≤16 tuổi): ")
            _set_font(r1, 12)
            r2 = p.add_run("." * 35 + " CCCD của mẹ hoặc người giám hộ: " + "." * 15)
            _set_font(r2, 12)
        elif "Mối quan hệ với trẻ:" in txt:
            p.text = ""
            r1 = p.add_run(
                f"13. Mối quan hệ với trẻ:\t{CHECK} Cha\t{CHECK} Mẹ\t{CHECK} Ông/bà\t"
                f"{CHECK} Anh/chị\t{CHECK} Họ hàng\t{CHECK} Khác"
            )
            _set_font(r1, 12)
        elif "Điện thoại di động:" in txt:
            p.text = ""
            r1 = p.add_run("14. Điện thoại di động: ")
            _set_font(r1, 12)
            r2 = p.add_run(str(rec.get("so_dien_thoai") or ""))
            _set_font(r2, 12, bold=True)
        elif "Lý do khám sức khỏe:" in txt:
            p.text = ""
            r1 = p.add_run("15. Lý do khám sức khỏe: ")
            _set_font(r1, 12)
            r2 = p.add_run("Khám sức khỏe học sinh định kỳ")
            _set_font(r2, 12)

    return d


def _build_form_scratch(rec: dict, group: dict, hospital_name: str = "BỆNH VIỆN HỒNG ĐỨC II") -> "Document":
    d = Document()
    section = d.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # style mặc định
    normal = d.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(12)

    gioi = (rec.get("gioi_tinh") or "").strip()

    # 1. Dòng Mẫu 2 trên cùng (Font size 12.5pt, in đậm, căn giữa như Mau02.pdf)
    _p(d, "Mẫu 2. MẪU GIẤY KHÁM SỨC KHỎE VÀ KHÁM SỨC KHỎE ĐỊNH KỲ DÙNG CHO TRẺ TỪ ĐỦ 06 TUỔI ĐẾN 18 TUỔI", size=12.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _p(d, "-----", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # 2. Header Table 2 cột (Trái: Logo + Tên BV + Số / Phải: Quốc hiệu)
    logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "images", "logo_hongduc2.png")
    tbl_hdr = d.add_table(rows=1, cols=2)
    tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_l, cell_r = tbl_hdr.rows[0].cells
    cell_l.width = Cm(6.5)
    cell_r.width = Cm(11.5)

    p_logo = cell_l.paragraphs[0]
    p_logo.paragraph_format.space_after = Pt(1)
    if os.path.exists(logo_path):
        r_logo = p_logo.add_run()
        r_logo.add_picture(logo_path, width=Cm(4.8))

    p_hname = cell_l.add_paragraph()
    p_hname.paragraph_format.space_after = Pt(2)
    _set_font(p_hname.add_run(hospital_name.upper()), 10, bold=True)

    p_so = cell_l.add_paragraph()
    p_so.paragraph_format.space_after = Pt(2)
    _set_font(p_so.add_run("Số: ……./GKSK-........."), 11)

    p_r1 = cell_r.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r1.paragraph_format.space_after = Pt(2)
    _set_font(p_r1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"), 12, bold=True)

    p_r2 = cell_r.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.paragraph_format.space_after = Pt(2)
    _set_font(p_r2.add_run("Độc lập - Tự do - Hạnh phúc"), 11, bold=True)

    p_r3 = cell_r.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r3.paragraph_format.space_after = Pt(4)
    _set_font(p_r3.add_run("_________________________"), 10)

    # 3. Tiêu đề MẪU 02 nằm ở giữa phía dưới "Số: ...."
    _p(d, "MẪU 02 - GIẤY KHÁM SỨC KHỎE VÀ KHÁM SỨC KHỎE ĐỊNH KỲ DÙNG CHO TRẺ TỪ ĐỦ 06 TUỔI ĐẾN DƯỚI 18 TUỔI", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
    _p(d, "(Ban hành kèm theo Thông tư số 25/2026/TT-BYT ngày 30 tháng 6 năm 2026 của Bộ Y tế)", size=10, italic=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # ---------- THÔNG TIN HÀNH CHÍNH (Đầy đủ 15 mục theo Mẫu 02) ----------
    _heading(d, "THÔNG TIN HÀNH CHÍNH")
    _field_line(d, "1. Họ và tên (viết chữ in hoa)", (rec.get("ho_ten") or "").upper())
    _checkbox_line(d, [("Nam", gioi == "Nam"), ("Nữ", gioi == "Nữ")])
    _field_line(d, "3. Ngày tháng năm sinh", rec.get("ngay_sinh"))
    _field_line(d, "4. Dân tộc", rec.get("dan_toc"))
    _field_line(d, "5. Nhóm máu (nếu có)", "")
    _add_boxes_line(d, "6. Số CCCD/Mã số định danh/Hộ chiếu", rec.get("cccd"), 12, label_cm=8.5)
    _add_boxes_line(d, "7. Số thẻ BHYT", rec.get("ma_bhyt"), 15, label_cm=8.5)
    dia_chi = ", ".join(x for x in [rec.get("so_nha"), rec.get("khu_pho"), rec.get("phuong"), rec.get("tinh")] if x)
    _field_line(d, "8. Nơi ở hiện tại", dia_chi)
    _field_line(d, "Xã/Phường", rec.get("phuong"))

    # 9, 10, 11. Thông tin trường học
    ten_truong = group.get("ten_doan") or ""
    _checkbox_line(d, [("Có", True), ("Không (chuyển qua câu 12)", False)], prefix="9. Trẻ có đi học: ")
    _field_line(d, "10. Tên Trường (nếu có)", ten_truong)
    _field_line(d, "11. Địa chỉ trường", group.get("dia_diem") or "")
    _field_line(d, "Xã/Phường", "")

    # 12, 13, 14, 15. Người giám hộ & liên hệ
    _field_line(d, "12. Họ và tên mẹ hoặc người giám hộ (đối với trẻ ≤16 tuổi)", "")
    _field_line(d, "CCCD của mẹ hoặc người giám hộ", "")
    _checkbox_line(d, [("Cha", False), ("Mẹ", False), ("Ông/bà", False), ("Anh/chị", False), ("Họ hàng", False), ("Khác", False)], prefix="13. Mối quan hệ với trẻ: ")
    _field_line(d, "14. Điện thoại di động", rec.get("so_dien_thoai"))
    _field_line(d, "15. Lý do khám sức khỏe", "Khám sức khỏe học sinh / định kỳ")

    _heading(d, "TIỀN SỬ BỆNH TẬT", size=12)
    _p(d, "1. Tiền sử gia đình (bao gồm bố mẹ, anh chị em ruột)", size=11, bold=True)
    _p(d, "Có ai trong gia đình mắc bệnh bẩm sinh hoặc bệnh truyền nhiễm không:", size=11)
    _checkbox_line(d, [("Không", False), ("Có. Nếu có, ghi cụ thể tên bệnh: " + "." * 40, False)])
    _p(d, "2. Tiền sử bản thân:", size=11, bold=True)
    _p(d, "a) Sản khoa:", size=11)
    _checkbox_line(d, [("Bình thường", False)])
    _checkbox_line(d, [("Không bình thường: Đẻ thiếu tháng", False), ("Đẻ ngạt", False)])
    _checkbox_line(d, [("Đẻ thừa cân", False), ("Đẻ có can thiệp", False)])
    _p(d, "b) Tiền sử bệnh/tật (bệnh bẩm sinh và mãn tính):", size=11)
    _checkbox_line(d, [("Không", False), ("Có, ghi cụ thể tên bệnh: " + "." * 40, False)])
    _p(d, "c) Hiện đang điều trị bệnh gì không:", size=11)
    _checkbox_line(d, [("Không", False), ("Có, ghi rõ tên bệnh và thuốc đang dùng: " + "." * 30, False)])

    _p(d, "Tôi xin cam đoan những điều khai trên đây hoàn toàn đúng với sự thật.", size=10.5, italic=True,
       space_before=8)
    _signature_block(d, "Người đề nghị khám sức khỏe")
    _hr(d)

    # ---------- ĐỐI TƯỢNG - CHI TRẢ ----------
    _heading(d, "THÔNG TIN ĐỐI TƯỢNG - CHI TRẢ")
    _field_line(d, "Đối tượng khám", group.get("doi_tuong_kham"))
    _field_line(d, "Hình thức chi trả", group.get("hinh_thuc_chi_tra"))
    _field_line(d, "Địa điểm khám", group.get("dia_diem"))
    _field_line(d, "Đoàn khám", group.get("ten_doan"))

    # ---------- KHÁM THỂ LỰC ----------
    _heading(d, "I. KHÁM THỂ LỰC")
    tbody = _new_table(d, 3, 2, widths=[9, 9])
    rows_txt = [
        ("- Chiều cao: …………… cm", "Mạch: …………… lần/phút"),
        ("- Cân nặng: …………… kg", "Huyết áp: …………… mmHg"),
        ("- Chỉ số BMI: ……………", "Nhịp thở: …………… lần/phút"),
    ]
    for i, (a, b) in enumerate(rows_txt):
        _set_cell_text(tbody.rows[i].cells[0], a, size=11)
        _set_cell_text(tbody.rows[i].cells[1], b, size=11)
    p = _p(d, "Phân loại thể lực: ", size=11)
    for loai in LOAI:
        _checkbox_run(p, f"Loại {loai}", size=11)

    # ---------- KHÁM LÂM SÀNG ----------
    _heading(d, "II. KHÁM LÂM SÀNG")
    for num, name, subs in CLINICAL_SPECIALTIES:
        _p(d, f"{num} {name}", size=11.5, bold=True)
        for sub_num, sub_name in subs:
            _p(d, f"{sub_num} {sub_name}", size=11)
            p = d.paragraphs[-1]
            _checkbox_run(p, "Chưa phát hiện bất thường", size=10.5)
            _field_line(d, "     Chẩn đoán sơ bộ (ICD)", "", dots=45)
            _field_line(d, "     Chẩn đoán xác định (ICD)", "", dots=45)
            p2 = _p(d, "     Phân loại: ", size=10.5)
            for loai in LOAI:
                _checkbox_run(p2, loai, size=10.5)
    for num, name in OTHER_SPECIALTIES:
        _p(d, f"{num} {name}", size=11.5, bold=True)
        if name == "Mắt":
            _field_line(d, "     Thị lực không kính: Mắt phải ... / Mắt trái", "", dots=30)
            _field_line(d, "     Thị lực có kính: Mắt phải ... / Mắt trái", "", dots=30)
        if name == "Tai - Mũi - Họng":
            _field_line(d, "     Thính lực: Tai trái (nói thường/thầm)", "", dots=30)
            _field_line(d, "     Thính lực: Tai phải (nói thường/thầm)", "", dots=30)
        if name == "Răng - Hàm - Mặt":
            _p(d, "     Sơ đồ răng: 18 17 16 15 14 13 12 11 | 21 22 23 24 25 26 27 28", size=9.5)
            _p(d, "                       48 47 46 45 44 43 42 41 | 31 32 33 34 35 36 37 38", size=9.5)
        p = d.paragraphs[-1]
        _checkbox_run(p, "Chưa phát hiện bất thường", size=10.5)
        _field_line(d, "     Chẩn đoán sơ bộ (ICD)", "", dots=45)
        _field_line(d, "     Chẩn đoán xác định (ICD)", "", dots=45)
        p2 = _p(d, "     Phân loại: ", size=10.5)
        for loai in LOAI:
            _checkbox_run(p2, loai, size=10.5)

    # ---------- CẬN LÂM SÀNG ----------
    _heading(d, "III. CẬN LÂM SÀNG")
    _field_line(d, "Xét nghiệm máu", "", dots=60)
    _field_line(d, "Xét nghiệm nước tiểu", "", dots=60)
    _field_line(d, "Chẩn đoán hình ảnh (X-quang, siêu âm...)", "", dots=55)
    _field_line(d, "Khác", "", dots=65)

    # ---------- KẾT LUẬN ----------
    _heading(d, "IV. KẾT LUẬN")
    p = _p(d, "1. Tình trạng sức khỏe: ", size=11.5)
    _checkbox_run(p, "Chưa phát hiện bất thường", size=11.5)
    _field_line(d, "Chẩn đoán", "", dots=60)
    p2 = _p(d, "2. Phân loại sức khỏe: ", size=11.5)
    for loai in LOAI:
        _checkbox_run(p2, f"Loại {loai}", size=11.5)
    _field_line(d, "3. Đề nghị", "", dots=65)
    _signature_block(d, "NGƯỜI KẾT LUẬN")

    # footer nhỏ để tra cứu
    foot = d.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = foot.add_run(f"Mã BN: {rec.get('id', '')}")
    _set_font(fr, 8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return d


# ---------------- xuất file / convert PDF ----------------
def save_docx(doc: "Document") -> bytes:
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_to_pdf(docx_bytes: bytes, timeout: int = 40, rec: dict = None, group: dict = None) -> bytes:
    """Convert .docx -> .pdf bằng LibreOffice headless (đảm bảo PDF khớp bản Word).
    Nếu LibreOffice chưa cài đặt hoặc gặp lỗi, tự động fallback sang ReportLab generator."""
    if hasattr(docx_bytes, "save"):
        docx_bytes = save_docx(docx_bytes)

    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "form.docx")
        with open(src, "wb") as f:
            f.write(docx_bytes)
        profile_dir = os.path.join(tmp, "lo_profile")
        os.makedirs(profile_dir, exist_ok=True)
        cmd = [
            "soffice", "--headless", "--norestore", "--nologo", "--nofirststartwizard",
            f"-env:UserInstallation=file://{profile_dir}",
            "--convert-to", "pdf", "--outdir", tmp, src,
        ]
        try:
            res = subprocess.run(cmd, capture_output=True, timeout=timeout)
            out = os.path.join(tmp, "form.pdf")
            if res.returncode == 0 and os.path.exists(out):
                with open(out, "rb") as f:
                    return f.read()
        except Exception:
            pass

    # Fallback to reportlab pdf_forms if soffice is not available or failed
    if rec is not None:
        from . import pdf_forms
        return pdf_forms.build_patient_form(rec, group or {})

    raise RuntimeError("Chưa cài LibreOffice (soffice) trong container và không có dữ liệu fallback.")


def merge_pdfs(pdf_bytes_list: list) -> bytes:
    import io
    from pypdf import PdfWriter
    writer = PdfWriter()
    for b in pdf_bytes_list:
        writer.append(io.BytesIO(b))
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def merge_docs(docs: list) -> "Document":
    """Gộp nhiều Document thành 1 (mỗi phiếu cách nhau bằng ngắt trang)."""
    if not docs:
        raise ValueError("Không có phiếu nào để gộp")
    base = docs[0]
    for extra in docs[1:]:
        base.add_page_break()
        for el in extra.element.body:
            base.element.body.append(el)
    return base
